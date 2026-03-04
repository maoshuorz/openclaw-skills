#!/usr/bin/env python3
"""
AI 智能文档生成器 - 合同与单据自动填充/批量生成

功能:
1. 模板解析: 从 .docx 模板中提取占位符字段
2. AI 智能填充: 根据用户模糊描述，AI 自动推断并填充字段
3. 批量生成: 根据简单要求批量制作多份文档
4. AI 自由生成: 无需模板，根据描述直接生成美观的合同/单据
"""

import json
import os
import re
import sys
import copy
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    Document = None

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None


class AIDocGenerator:
    """AI 智能文档生成器"""

    # 模板中的占位符格式: {{字段名}} 或 {字段名}
    PLACEHOLDER_PATTERN = re.compile(r'\{\{(.+?)\}\}|\{([A-Za-z\u4e00-\u9fff_]+)\}')

    def __init__(self, api_key, base_url=None, model="gpt-4o"):
        """
        初始化生成器

        Args:
            api_key: OpenAI 兼容 API 的密钥
            base_url: API 基础地址 (可选, 用于兼容其他服务)
            model: 使用的模型名称
        """
        if OpenAI is None:
            raise ImportError("请安装 openai: pip install openai")
        if Document is None:
            raise ImportError("请安装 python-docx: pip install python-docx")

        client_kwargs = {"api_key": api_key}
        if base_url:
            client_kwargs["base_url"] = base_url
        self.client = OpenAI(**client_kwargs)
        self.model = model

    # ─── 模板解析 ───────────────────────────────────────────

    def parse_template(self, template_path):
        """
        解析 .docx 模板, 提取所有占位符字段

        Args:
            template_path: 模板文件路径

        Returns:
            dict: {fields: [字段列表], template_path: 路径}
        """
        doc = Document(template_path)
        fields = set()

        # 从段落中提取
        for para in doc.paragraphs:
            for match in self.PLACEHOLDER_PATTERN.finditer(para.text):
                fields.add(match.group(1) or match.group(2))

        # 从表格中提取
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for match in self.PLACEHOLDER_PATTERN.finditer(cell.text):
                        fields.add(match.group(1) or match.group(2))

        # 从页眉页脚中提取
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                for para in header_footer.paragraphs:
                    for match in self.PLACEHOLDER_PATTERN.finditer(para.text):
                        fields.add(match.group(1) or match.group(2))

        return {
            "fields": sorted(fields),
            "template_path": template_path,
            "field_count": len(fields)
        }

    # ─── AI 智能字段提取 ─────────────────────────────────────

    def _ai_extract_fields(self, fields, user_description):
        """
        使用 AI 从模糊描述中提取结构化字段值

        Args:
            fields: 模板中的字段列表
            user_description: 用户的模糊描述

        Returns:
            dict: 字段名 -> 值 的映射
        """
        prompt = f"""你是一个专业的合同和单据填写助手。

用户提供了以下模糊描述，请根据描述推断出所有字段的值。
如果某个字段无法从描述中推断，请根据常理和上下文生成合理的默认值。
日期类字段如未指定，默认使用今天: {datetime.now().strftime('%Y年%m月%d日')}。
金额类字段请同时提供大写和小写（如适用）。

需要填写的字段:
{json.dumps(fields, ensure_ascii=False, indent=2)}

用户描述:
{user_description}

请以 JSON 格式返回，key 为字段名，value 为填入值。只返回 JSON，不要其他内容。"""

        response = self.client.chat.completions.create(
            model=self.model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
        )

        text = response.choices[0].message.content.strip()
        # 提取 JSON 块
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', text)
        if json_match:
            text = json_match.group(1)
        return json.loads(text)

    # ─── 模板填充 ─────────────────────────────────────────

    def _replace_in_paragraph(self, paragraph, field_values):
        """在段落中替换占位符，保留原始格式"""
        full_text = paragraph.text
        if not self.PLACEHOLDER_PATTERN.search(full_text):
            return

        new_text = full_text
        for field, value in field_values.items():
            new_text = new_text.replace(f'{{{{{field}}}}}', str(value))
            new_text = new_text.replace(f'{{{field}}}', str(value))

        if new_text != full_text:
            # 保留第一个 run 的格式，替换文本
            if paragraph.runs:
                first_run_font = paragraph.runs[0].font
                for i, run in enumerate(paragraph.runs):
                    if i == 0:
                        continue
                    run.text = ""
                paragraph.runs[0].text = new_text
            else:
                paragraph.text = new_text

    def fill_template(self, template_path, field_values, output_path):
        """
        填充模板并保存

        Args:
            template_path: 模板文件路径
            field_values: 字段值字典 {字段名: 值}
            output_path: 输出文件路径

        Returns:
            dict: 结果信息
        """
        doc = Document(template_path)

        # 替换段落
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, field_values)

        # 替换表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, field_values)

        # 替换页眉页脚
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                for para in header_footer.paragraphs:
                    self._replace_in_paragraph(para, field_values)

        doc.save(output_path)
        return {"output": output_path, "fields_filled": len(field_values)}

    # ─── AI 智能填充 (模糊信息 → 填充模板) ─────────────────

    def smart_fill(self, template_path, user_description, output_path):
        """
        智能填充: 用户提供模糊描述，AI 自动推断字段值并填充模板

        Args:
            template_path: 模板文件路径
            user_description: 用户的模糊描述
            output_path: 输出文件路径

        Returns:
            dict: {output, fields, values}
        """
        # 1. 解析模板获取字段
        parsed = self.parse_template(template_path)
        fields = parsed["fields"]

        # 2. AI 推断字段值
        field_values = self._ai_extract_fields(fields, user_description)

        # 3. 填充模板
        result = self.fill_template(template_path, field_values, output_path)
        result["fields"] = fields
        result["values"] = field_values
        return result

    # ─── 批量生成 ──────────────────────────────────────────

    def _ai_generate_batch_data(self, fields, user_description, count):
        """使用 AI 批量生成多组数据"""
        prompt = f"""你是一个专业的合同和单据填写助手。

用户需要批量生成 {count} 份文档。请根据用户描述，为每份文档生成一组合理的字段值。
每组数据应当有合理的差异（如不同的名称、日期、编号等）。
日期类字段默认从今天 {datetime.now().strftime('%Y年%m月%d日')} 开始。
编号类字段请生成连续编号。

需要填写的字段:
{json.dumps(fields, ensure_ascii=False, indent=2)}

用户描述:
{user_description}

请以 JSON 数组格式返回 {count} 组数据，每组为一个对象。只返回 JSON，不要其他内容。"""

        response = self.client.chat.completions.create(
            model=self.model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )

        text = response.choices[0].message.content.strip()
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', text)
        if json_match:
            text = json_match.group(1)
        return json.loads(text)

    def batch_fill(self, template_path, user_description, output_dir, count=5):
        """
        批量填充: 根据描述批量生成多份文档

        Args:
            template_path: 模板文件路径
            user_description: 用户对批量文档的描述
            output_dir: 输出目录
            count: 生成数量

        Returns:
            dict: {outputs: [文件列表], count: 数量}
        """
        os.makedirs(output_dir, exist_ok=True)

        # 1. 解析模板
        parsed = self.parse_template(template_path)
        fields = parsed["fields"]

        # 2. AI 批量生成数据
        batch_data = self._ai_generate_batch_data(fields, user_description, count)

        # 3. 逐份填充
        outputs = []
        base_name = os.path.splitext(os.path.basename(template_path))[0]
        for i, field_values in enumerate(batch_data):
            output_path = os.path.join(output_dir, f"{base_name}_{i+1:03d}.docx")
            self.fill_template(template_path, field_values, output_path)
            outputs.append({"file": output_path, "values": field_values})

        return {"outputs": outputs, "count": len(outputs)}

    # ─── AI 自由生成 (无需模板) ────────────────────────────

    def ai_generate(self, user_description, output_path, doc_type="合同"):
        """
        AI 自由生成: 无需模板，根据描述直接生成完整文档

        Args:
            user_description: 用户对文档的描述
            output_path: 输出文件路径
            doc_type: 文档类型 (合同/收据/发票/协议/报价单 等)

        Returns:
            dict: {output, doc_type, sections}
        """
        prompt = f"""你是一个专业的{doc_type}撰写专家。

请根据用户描述，生成一份完整、专业、美观的{doc_type}。

要求:
1. 内容完整，条款严谨，符合法律规范
2. 结构清晰，格式规范
3. 如果用户描述模糊，请根据常理补全所有必要信息
4. 日期默认今天: {datetime.now().strftime('%Y年%m月%d日')}

用户描述:
{user_description}

请以以下 JSON 格式返回文档内容:
{{
  "title": "文档标题",
  "subtitle": "副标题(可选，没有则为空字符串)",
  "sections": [
    {{
      "heading": "章节标题(可选，没有则为空字符串)",
      "content": "章节内容（支持多行，用\\n换行）"
    }}
  ],
  "table": {{
    "headers": ["列1", "列2"],
    "rows": [["值1", "值2"]]
  }},
  "signatures": [
    {{
      "party": "签署方名称",
      "fields": ["签字:", "日期:", "盖章:"]
    }}
  ],
  "footer": "页脚文字(可选)"
}}

注意:
- table 为可选，没有则设为 null
- sections 中可以有多个章节
- 只返回 JSON，不要其他内容"""

        response = self.client.chat.completions.create(
            model=self.model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
        )

        text = response.choices[0].message.content.strip()
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', text)
        if json_match:
            text = json_match.group(1)
        doc_data = json.loads(text)

        # 构建 Word 文档
        doc = Document()

        # ── 页面设置 ──
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

        # ── 标题 ──
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(doc_data.get("title", doc_type))
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.name = "SimHei"
        title_para.space_after = Pt(6)

        # ── 副标题 ──
        subtitle = doc_data.get("subtitle", "")
        if subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.size = Pt(12)
            sub_run.font.color.rgb = RGBColor(100, 100, 100)
            sub_para.space_after = Pt(12)

        # ── 正文章节 ──
        for sec in doc_data.get("sections", []):
            heading = sec.get("heading", "")
            content = sec.get("content", "")

            if heading:
                h_para = doc.add_paragraph()
                h_run = h_para.add_run(heading)
                h_run.font.size = Pt(14)
                h_run.font.bold = True
                h_para.space_before = Pt(12)
                h_para.space_after = Pt(4)

            if content:
                for line in content.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.size = Pt(11)
                    p.paragraph_format.first_line_indent = Cm(0.74)
                    p.paragraph_format.line_spacing = 1.5

        # ── 表格 ──
        table_data = doc_data.get("table")
        if table_data and table_data.get("headers"):
            doc.add_paragraph()  # 空行
            headers = table_data["headers"]
            rows = table_data.get("rows", [])
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 表头
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

            # 数据行
            for r_idx, row in enumerate(rows):
                for c_idx, val in enumerate(row):
                    cell = table.rows[r_idx + 1].cells[c_idx]
                    cell.text = str(val)
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(10)

        # ── 签署区 ──
        signatures = doc_data.get("signatures", [])
        if signatures:
            doc.add_paragraph()  # 空行
            sig_para = doc.add_paragraph()
            sig_para.space_before = Pt(24)

            for sig in signatures:
                party = sig.get("party", "")
                fields = sig.get("fields", [])
                p = doc.add_paragraph()
                run = p.add_run(party)
                run.font.bold = True
                run.font.size = Pt(11)
                for f in fields:
                    fp = doc.add_paragraph()
                    frun = fp.add_run(f)
                    frun.font.size = Pt(11)
                    fp.paragraph_format.left_indent = Cm(0.5)
                doc.add_paragraph()  # 签署方之间空行

        # ── 页脚 ──
        footer_text = doc_data.get("footer", "")
        if footer_text:
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            frun = fp.add_run(footer_text)
            frun.font.size = Pt(8)
            frun.font.color.rgb = RGBColor(150, 150, 150)

        doc.save(output_path)
        return {
            "output": output_path,
            "doc_type": doc_type,
            "title": doc_data.get("title", ""),
            "sections": len(doc_data.get("sections", []))
        }

    def ai_batch_generate(self, user_description, output_dir, count=5, doc_type="合同"):
        """
        AI 批量自由生成: 批量生成多份不同内容的文档

        Args:
            user_description: 用户描述
            output_dir: 输出目录
            count: 生成数量
            doc_type: 文档类型

        Returns:
            dict: {outputs: [文件列表], count: 数量}
        """
        os.makedirs(output_dir, exist_ok=True)
        outputs = []
        for i in range(count):
            output_path = os.path.join(output_dir, f"{doc_type}_{i+1:03d}.docx")
            desc = f"{user_description}\n\n这是第 {i+1}/{count} 份，请确保内容与其他份有合理差异（如编号、名称等不同）。当前编号: {i+1:03d}"
            result = self.ai_generate(desc, output_path, doc_type)
            outputs.append(result)
        return {"outputs": outputs, "count": len(outputs)}


# ─── CLI 入口 ────────────────────────────────────────────

def print_usage():
    print("""AI 智能文档生成器

用法:
  python ai_doc_generator.py <command> [options]

命令:
  parse     <template.docx>                          解析模板，列出所有占位符字段
  fill      <template.docx> <fields.json> <out.docx> 使用 JSON 数据填充模板
  smart     <template.docx> <description> <out.docx> AI 智能填充 (模糊描述自动填充)
  batch     <template.docx> <description> <out_dir> <count>  批量生成
  generate  <description> <out.docx> [doc_type]      AI 自由生成文档
  batch-gen <description> <out_dir> <count> [doc_type]       批量自由生成

环境变量:
  OPENAI_API_KEY   - API 密钥 (必需)
  OPENAI_BASE_URL  - API 地址 (可选, 兼容其他服务)
  OPENAI_MODEL     - 模型名称 (可选, 默认 gpt-4o)""")


if __name__ == '__main__':
    args = sys.argv[1:]
    if not args:
        print_usage()
        sys.exit(1)

    cmd = args[0]

    if cmd == "parse":
        if len(args) < 2:
            print("用法: python ai_doc_generator.py parse <template.docx>")
            sys.exit(1)
        gen = AIDocGenerator.__new__(AIDocGenerator)
        gen.PLACEHOLDER_PATTERN = AIDocGenerator.PLACEHOLDER_PATTERN
        result = gen.parse_template(args[1])
        print(json.dumps(result, ensure_ascii=False, indent=2))

    elif cmd == "fill":
        if len(args) < 4:
            print("用法: python ai_doc_generator.py fill <template.docx> <fields.json> <out.docx>")
            sys.exit(1)
        gen = AIDocGenerator.__new__(AIDocGenerator)
        gen.PLACEHOLDER_PATTERN = AIDocGenerator.PLACEHOLDER_PATTERN
        with open(args[2], 'r', encoding='utf-8') as f:
            field_values = json.load(f)
        result = gen.fill_template(args[1], field_values, args[3])
        print(json.dumps(result, ensure_ascii=False, indent=2))

    elif cmd in ("smart", "batch", "generate", "batch-gen"):
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            print("错误: 请设置环境变量 OPENAI_API_KEY")
            sys.exit(1)
        base_url = os.environ.get("OPENAI_BASE_URL")
        model = os.environ.get("OPENAI_MODEL", "gpt-4o")
        gen = AIDocGenerator(api_key, base_url, model)

        if cmd == "smart":
            if len(args) < 4:
                print("用法: python ai_doc_generator.py smart <template.docx> <description> <out.docx>")
                sys.exit(1)
            result = gen.smart_fill(args[1], args[2], args[3])

        elif cmd == "batch":
            if len(args) < 5:
                print("用法: python ai_doc_generator.py batch <template.docx> <description> <out_dir> <count>")
                sys.exit(1)
            result = gen.batch_fill(args[1], args[2], args[3], int(args[4]))

        elif cmd == "generate":
            if len(args) < 3:
                print("用法: python ai_doc_generator.py generate <description> <out.docx> [doc_type]")
                sys.exit(1)
            doc_type = args[3] if len(args) > 3 else "合同"
            result = gen.ai_generate(args[1], args[2], doc_type)

        elif cmd == "batch-gen":
            if len(args) < 4:
                print("用法: python ai_doc_generator.py batch-gen <description> <out_dir> <count> [doc_type]")
                sys.exit(1)
            doc_type = args[4] if len(args) > 4 else "合同"
            result = gen.ai_batch_generate(args[1], args[2], int(args[3]), doc_type)

        print(json.dumps(result, ensure_ascii=False, indent=2))

    else:
        print(f"未知命令: {cmd}")
        print_usage()
        sys.exit(1)
