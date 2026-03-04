#!/usr/bin/env python3
"""Demo: 生成一份软件开发外包合同，展示效果"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 页面设置
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ── 标题 ──
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("软件开发外包服务合同")
title_run.font.size = Pt(22)
title_run.font.bold = True
title_para.space_after = Pt(6)

# ── 副标题 ──
sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_para.add_run("合同编号: HT-2026-0304-001")
sub_run.font.size = Pt(12)
sub_run.font.color.rgb = RGBColor(100, 100, 100)
sub_para.space_after = Pt(12)

# ── 甲乙方信息 ──
party_lines = [
    "甲方（委托方）: 北京星云科技有限公司",
    "统一社会信用代码: 91110108MA01XXXXX",
    "法定代表人: 张明远",
    "地址: 北京市海淀区中关村科技园区创新大厦12层",
    "",
    "乙方（受托方）: 深圳锐码软件工作室",
    "统一社会信用代码: 91440300MA5FXXXXX",
    "负责人: 李文博",
    "地址: 广东省深圳市南山区科技园南区数码大厦8层806室",
]

for line in party_lines:
    if not line:
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.size = Pt(11)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5

# ── 条款 ──
clauses = [
    ("第一条 项目概况",
     "甲方委托乙方进行「星云电商App」移动应用软件的设计与开发工作（以下简称"
     "本项目）。本项目包括但不限于需求分析、UI/UX设计、前端开发、后端开发、"
     "接口对接、系统测试及部署上线等全部工作内容。"),

    ("第二条 合同金额与支付方式",
     "本合同总金额为人民币伍拾万元整（\\u00a5500,000.00）。付款分三期进行：\n\n"
     "第一期：合同签订后5个工作日内，甲方支付合同总额的30%，即人民币壹拾伍万元整"
     "（\\u00a5150,000.00），作为项目启动款。\n\n"
     "第二期：项目开发完成并通过甲方初步验收后5个工作日内，甲方支付合同总额的50%，"
     "即人民币贰拾伍万元整（\\u00a5250,000.00）。\n\n"
     "第三期：项目最终验收合格并上线运行满30日后5个工作日内，甲方支付合同总额的20%，"
     "即人民币壹拾万元整（\\u00a5100,000.00），作为尾款。"),

    ("第三条 开发周期与里程碑",
     "本项目开发周期为120个自然日，自合同签订之日起计算。具体里程碑安排见下表："),

    ("TABLE", None),  # 占位，后面插表格

    ("第四条 知识产权",
     "本项目开发完成并甲方付清全部款项后，项目全部源代码、设计文件、文档资料等知识产权"
     "归甲方所有。乙方不得将本项目的代码、设计及相关技术资料用于其他商业用途或转让给第三方。"),

    ("第五条 保密条款",
     "双方应对在合作过程中知悉的对方商业秘密、技术秘密及其他保密信息严格保密。保密期限自"
     "合同签订之日起三年内有效。违反保密义务的一方应赔偿对方因此遭受的全部损失。"),

    ("第六条 质量保证与维护",
     "乙方应保证交付的软件产品质量符合甲方需求文档的要求，不存在严重缺陷。项目验收合格后，"
     "乙方提供12个月的免费维护期，包括Bug修复及小幅功能调整。维护期内响应时间不超过24小时。"),

    ("第七条 违约责任",
     "甲方逾期付款的，每逾期一日应向乙方支付逾期金额千分之一的违约金。乙方逾期交付的，"
     "每逾期一日应向甲方支付合同总额千分之一的违约金。任何一方违约导致合同无法继续履行的，"
     "违约方应赔偿守约方的直接经济损失。"),

    ("第八条 争议解决",
     "因本合同引起的或与本合同有关的任何争议，双方应首先通过友好协商解决。协商不成的，"
     "任何一方均可向甲方所在地有管辖权的人民法院提起诉讼。"),

    ("第九条 其他",
     "本合同一式两份，甲乙双方各执一份，自双方签字盖章之日起生效，具有同等法律效力。"
     "本合同未尽事宜，双方可另行签订补充协议，补充协议与本合同具有同等法律效力。"),
]

for heading, content in clauses:
    if heading == "TABLE":
        # 里程碑表格
        doc.add_paragraph()
        headers = ["阶段", "内容", "时间节点", "交付物"]
        rows = [
            ["第一阶段", "需求分析与UI设计", "第1-30日", "需求文档、设计稿"],
            ["第二阶段", "前端与后端核心开发", "第31-75日", "开发版本Demo"],
            ["第三阶段", "联调测试与优化", "第76-100日", "测试报告"],
            ["第四阶段", "部署上线与验收", "第101-120日", "上线版本、源代码"],
        ]
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = val
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(10)
        doc.add_paragraph()
        continue

    # 章节标题
    h_para = doc.add_paragraph()
    h_run = h_para.add_run(heading)
    h_run.font.size = Pt(14)
    h_run.font.bold = True
    h_para.space_before = Pt(12)
    h_para.space_after = Pt(4)

    # 章节内容
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5

# ── 签署区 ──
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("甲方（盖章）: 北京星云科技有限公司")
run.font.bold = True
run.font.size = Pt(11)
for field in ["法定代表人/授权代表签字:                    ",
              "日期:        年    月    日"]:
    fp = doc.add_paragraph()
    frun = fp.add_run(field)
    frun.font.size = Pt(11)
    fp.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("乙方（盖章）: 深圳锐码软件工作室")
run.font.bold = True
run.font.size = Pt(11)
for field in ["负责人/授权代表签字:                    ",
              "日期:        年    月    日"]:
    fp = doc.add_paragraph()
    frun = fp.add_run(field)
    frun.font.size = Pt(11)
    fp.paragraph_format.left_indent = Cm(0.5)

# ── 页脚 ──
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
frun = fp.add_run("本合同为机密文件，未经双方书面同意不得向第三方披露")
frun.font.size = Pt(8)
frun.font.color.rgb = RGBColor(150, 150, 150)

output_path = "/home/user/openclaw-skills/ai-doc-generator/demo_contract.docx"
doc.save(output_path)
print(f"Done: {output_path}")
